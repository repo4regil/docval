"""
footer_validator.py
-------------------
Validates the content of the document footer section.

Three checks are implemented:
  1. Confidentiality label (text search)
  2. Page number field (XML field detection)
  3. Company name (configurable text search)

To add a new check, write a ``_check_<name>(footer, logger)`` function and
register it in the ``check_map`` at the bottom of ``validate_footer``.
"""

import logging
import re
from typing import List, Optional

from docx import Document

from config import FOOTER_CHECKS, COMPANY_NAME
from models import Status, ValidationResult

CATEGORY = "Footer"


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------

def validate_footer(
    doc: Document,
    logger: logging.Logger,
    total_pages: Optional[int] = None,
) -> List[ValidationResult]:
    """
    Run all enabled footer checks and return their results.

    Parameters
    ----------
    doc : docx.Document
        An opened python-docx Document.
    logger : logging.Logger
        Session logger.

    Returns
    -------
    list of ValidationResult
    """
    logger.info("--- Footer validation started ---")
    results: List[ValidationResult] = []

    try:
        footer = doc.sections[0].footer
    except IndexError:
        logger.error("Document has no sections – cannot read footer.")
        results.append(ValidationResult(
            category=CATEGORY,
            check="footer_access",
            label="Footer Access",
            status=Status.FAIL,
            message="Document has no sections; footer could not be read.",
        ))
        return results

    check_map = {
        "has_confidentiality_label": _check_confidentiality_label,
        "has_page_number_field":     _check_page_number_field,
        "has_company_name":          _check_company_name,
    }

    for check_key, check_fn in check_map.items():
        if not FOOTER_CHECKS.get(check_key, True):
            results.append(ValidationResult(
                category=CATEGORY,
                check=check_key,
                label=check_key.replace("_", " ").title(),
                status=Status.SKIP,
                message="Check disabled in configuration.",
            ))
            logger.debug("Skipping footer check: %s", check_key)
            continue

        if check_key == "has_page_number_field":
            result = check_fn(footer, logger, total_pages=total_pages)
        else:
            result = check_fn(footer, logger)
        results.append(result)
        logger.info("[%s] %s – %s", CATEGORY, result.check, result.status.value)

    logger.info("--- Footer validation complete ---")
    return results


# ---------------------------------------------------------------------------
# Individual check functions
# ---------------------------------------------------------------------------

def _check_confidentiality_label(footer, logger: logging.Logger) -> ValidationResult:
    """
    Check that the footer contains a confidentiality / classification label
    (e.g. "Confidential", "Internal", "Restricted").
    """
    label = "Confidentiality Label in Footer"
    keywords = ["confidential", "internal", "restricted", "proprietary", "private"]
    try:
        text = " ".join(p.text for p in footer.paragraphs).lower()
        found = next((kw for kw in keywords if kw in text), None)
        if found:
            return ValidationResult(
                category=CATEGORY,
                check="has_confidentiality_label",
                label=label,
                status=Status.PASS,
                message=f'Confidentiality label "{found.title()}" found in footer.',
            )
        else:
            return ValidationResult(
                category=CATEGORY,
                check="has_confidentiality_label",
                label=label,
                status=Status.FAIL,
                message=(
                    "No confidentiality label found in the footer. "
                    f"Expected one of: {', '.join(kw.title() for kw in keywords)}."
                ),
            )
    except Exception as exc:  # noqa: BLE001
        logger.exception("Error checking footer confidentiality label: %s", exc)
        return ValidationResult(
            category=CATEGORY,
            check="has_confidentiality_label",
            label=label,
            status=Status.WARN,
            message=f"Error reading footer: {exc}",
        )


def _check_page_number_field(
    footer,
    logger: logging.Logger,
    total_pages: Optional[int] = None,
) -> ValidationResult:
    """
    Check that the footer contains a Word page-number field.

    Word represents page numbers as ``<w:fldSimple w:instr=" PAGE "/>`` or
    ``<w:instrText>`` containing PAGE / NUMPAGES.  We inspect the raw XML.
    """
    label = "Page Number Field in Footer"
    try:
        footer_xml = footer._element.xml
        has_page_field = (
            '" PAGE "' in footer_xml
            or ">PAGE<" in footer_xml
            or "NUMPAGES" in footer_xml
            or 'w:instrText' in footer_xml and "PAGE" in footer_xml
        )
        if not has_page_field:
            return ValidationResult(
                category=CATEGORY,
                check="has_page_number_field",
                label=label,
                status=Status.FAIL,
                message=(
                    "No page number field found in the footer. "
                    "Insert a PAGE field via Insert > Header & Footer > Page Number."
                ),
            )

        # If we don't have a computed total, we can only assert that the field exists.
        if total_pages is None:
            return ValidationResult(
                category=CATEGORY,
                check="has_page_number_field",
                label=label,
                status=Status.PASS,
                message="Page number field (PAGE/NUMPAGES) detected in the footer.",
            )

        # Try to parse the total-pages value mentioned in the footer text (e.g. 'Page 1 of 10').
        footer_text = " ".join(p.text for p in footer.paragraphs)
        match = re.search(r"\bof\s+(\d+)\b", footer_text, flags=re.IGNORECASE)
        if not match:
            # Fall back to XML-based NUMPAGES pattern if visible text doesn't contain it
            match = re.search(r"NUMPAGES[^0-9]*(\d+)", footer_xml)

        if match:
            stated_total = int(match.group(1))
            if stated_total == total_pages:
                return ValidationResult(
                    category=CATEGORY,
                    check="has_page_number_field",
                    label=label,
                    status=Status.PASS,
                    message=(
                        f"Footer page count matches document total: {stated_total} page(s)."
                    ),
                )
            else:
                return ValidationResult(
                    category=CATEGORY,
                    check="has_page_number_field",
                    label=label,
                    status=Status.FAIL,
                    message=(
                        f"Footer states total pages as {stated_total}, "
                        f"but the document has {total_pages} page(s)."
                    ),
                )

        # If we cannot parse the stated total, just confirm the field exists.
        return ValidationResult(
            category=CATEGORY,
            check="has_page_number_field",
            label=label,
            status=Status.WARN,
            message=(
                "Page number field detected, but could not determine the "
                "total pages stated in the footer to compare with the document."
            ),
        )
    except Exception as exc:  # noqa: BLE001
        logger.exception("Error checking footer page field: %s", exc)
        return ValidationResult(
            category=CATEGORY,
            check="has_page_number_field",
            label=label,
            status=Status.WARN,
            message=f"Error inspecting footer XML: {exc}",
        )


def _check_company_name(footer, logger: logging.Logger) -> ValidationResult:
    """
    Check that the company name (from ``config.COMPANY_NAME``) appears in
    the footer text.
    """
    label = "Company Name in Footer"
    try:
        text = " ".join(p.text for p in footer.paragraphs)
        if COMPANY_NAME.lower() in text.lower():
            return ValidationResult(
                category=CATEGORY,
                check="has_company_name",
                label=label,
                status=Status.PASS,
                message=f'Company name "{COMPANY_NAME}" found in footer.',
            )
        else:
            return ValidationResult(
                category=CATEGORY,
                check="has_company_name",
                label=label,
                status=Status.FAIL,
                message=(
                    f'Company name "{COMPANY_NAME}" not found in footer. '
                    "Update COMPANY_NAME in config.py if the name has changed."
                ),
            )
    except Exception as exc:  # noqa: BLE001
        logger.exception("Error checking footer company name: %s", exc)
        return ValidationResult(
            category=CATEGORY,
            check="has_company_name",
            label=label,
            status=Status.WARN,
            message=f"Error reading footer text: {exc}",
        )
