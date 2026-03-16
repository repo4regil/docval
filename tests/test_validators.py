"""
tests/test_validators.py
------------------------
Smoke tests for the Document Validator modules.

These tests use a minimal synthetic Word document created in-memory
(no real .docm file needed) to verify:
  - DocumentProperties extraction
  - Heading detection
  - Table extraction
  - ValidationReport accumulation

Run with:
    pytest tests/ -v
"""

import pytest
import io
import os
import sys
import logging

# Ensure the project root is on the path when running from the tests/ folder
sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

from docx import Document as DocxDocument
from docx.shared import Pt

from models import Status, ValidationResult, ValidationReport, DocumentProperties
from document_properties import extract_properties
from heading_validator import validate_heading
from table_extractor import get_tables_by_heading
from table_validator import validate_tables


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------

@pytest.fixture
def null_logger() -> logging.Logger:
    """A logger that discards all output – keeps test output clean."""
    logger = logging.getLogger("test_null")
    logger.addHandler(logging.NullHandler())
    return logger


@pytest.fixture
def minimal_doc():
    """
    A minimal python-docx Document with:
    - Core properties set (title, author)
    - One Heading 1 paragraph
    - A table under a heading named "Change History"
    """
    doc = DocxDocument()

    # Core properties
    doc.core_properties.title  = "Test Document"
    doc.core_properties.author = "Test Author"

    # Heading 1
    doc.add_heading("Test Document Title", level=1)

    # Body paragraph
    doc.add_paragraph("This is a body paragraph.")

    # Heading 2 for table
    doc.add_heading("Change History", level=2)

    # Table with header row
    table = doc.add_table(rows=3, cols=3)
    headers = ["Date", "Author", "Description"]
    for col_idx, header_text in enumerate(headers):
        table.rows[0].cells[col_idx].text = header_text
    table.rows[1].cells[0].text = "2026-01-01"
    table.rows[1].cells[1].text = "Alice"
    table.rows[1].cells[2].text = "Initial version"
    table.rows[2].cells[0].text = "2026-02-15"
    table.rows[2].cells[1].text = "Bob"
    table.rows[2].cells[2].text = "Review updates"

    return doc


# ---------------------------------------------------------------------------
# DocumentProperties tests
# ---------------------------------------------------------------------------

class TestDocumentProperties:

    def test_title_extracted(self, minimal_doc, null_logger):
        props = extract_properties(minimal_doc, null_logger)
        assert props.title == "Test Document"

    def test_author_extracted(self, minimal_doc, null_logger):
        props = extract_properties(minimal_doc, null_logger)
        assert props.author == "Test Author"

    def test_word_count_positive(self, minimal_doc, null_logger):
        props = extract_properties(minimal_doc, null_logger)
        assert props.word_count > 0

    def test_section_count_at_least_one(self, minimal_doc, null_logger):
        props = extract_properties(minimal_doc, null_logger)
        assert props.section_count >= 1

    def test_to_display_dict_no_none_keys(self, minimal_doc, null_logger):
        props = extract_properties(minimal_doc, null_logger)
        display = props.to_display_dict()
        # All values should be strings (None converted to "—")
        for key, val in display.items():
            assert isinstance(val, str), f"Key '{key}' value is not a string: {val!r}"


# ---------------------------------------------------------------------------
# Heading validator tests
# ---------------------------------------------------------------------------

class TestHeadingValidator:

    def test_heading_1_found(self, minimal_doc, null_logger):
        results = validate_heading(minimal_doc, null_logger)
        h1_result = next((r for r in results if r.check == "has_heading_1"), None)
        assert h1_result is not None
        assert h1_result.status == Status.PASS

    def test_no_heading_1_fails(self, null_logger):
        doc = DocxDocument()
        doc.add_paragraph("Just a plain paragraph, no heading.")
        results = validate_heading(doc, null_logger)
        h1_result = next((r for r in results if r.check == "has_heading_1"), None)
        assert h1_result is not None
        assert h1_result.status == Status.FAIL


# ---------------------------------------------------------------------------
# Table extractor tests
# ---------------------------------------------------------------------------

class TestTableExtractor:

    def test_table_found_under_heading(self, minimal_doc, null_logger):
        index = get_tables_by_heading(minimal_doc, ["Change History"], null_logger)
        assert "Change History" in index
        assert len(index["Change History"]) == 2  # 2 data rows

    def test_table_column_keys(self, minimal_doc, null_logger):
        index = get_tables_by_heading(minimal_doc, ["Change History"], null_logger)
        rows = index["Change History"]
        assert rows[0].get("Date") == "2026-01-01"
        assert rows[0].get("Author") == "Alice"

    def test_missing_heading_returns_empty(self, minimal_doc, null_logger):
        index = get_tables_by_heading(minimal_doc, ["Non-Existent Heading"], null_logger)
        assert index["Non-Existent Heading"] == []


# ---------------------------------------------------------------------------
# ValidationReport tests
# ---------------------------------------------------------------------------

class TestValidationReport:

    def test_is_valid_with_no_failures(self):
        report = ValidationReport(filename="test.docm")
        report.add(ValidationResult("Cat", "chk1", "Check 1", Status.PASS, "ok"))
        report.add(ValidationResult("Cat", "chk2", "Check 2", Status.WARN, "warning"))
        assert report.is_valid is True

    def test_is_invalid_with_failure(self):
        report = ValidationReport(filename="test.docm")
        report.add(ValidationResult("Cat", "chk1", "Check 1", Status.FAIL, "failed"))
        assert report.is_valid is False

    def test_summary_string(self):
        report = ValidationReport(filename="test.docm")
        report.add(ValidationResult("Cat", "chk1", "A", Status.PASS, "ok"))
        report.add(ValidationResult("Cat", "chk2", "B", Status.FAIL, "no"))
        summary = report.summary()
        assert "passed" in summary
        assert "failed" in summary

    def test_by_category_groups_correctly(self):
        report = ValidationReport(filename="test.docm")
        report.add(ValidationResult("Header", "h1", "H1", Status.PASS, "ok"))
        report.add(ValidationResult("Footer", "f1", "F1", Status.FAIL, "no"))
        report.add(ValidationResult("Header", "h2", "H2", Status.PASS, "ok"))
        grouped = report.by_category()
        assert len(grouped["Header"]) == 2
        assert len(grouped["Footer"]) == 1
